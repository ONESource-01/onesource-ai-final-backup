import os
from contextlib import asynccontextmanager
from motor.motor_asyncio import AsyncIOMotorClient
from fastapi import FastAPI, APIRouter, Depends, HTTPException, Request, UploadFile, File, Form
from fastapi.security import HTTPBearer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, EmailStr, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timedelta
import openai
from openai import AsyncOpenAI
import io
import base64
import mimetypes
import hashlib
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import logging
from pathlib import Path
from dotenv import load_dotenv

# Import AI Intelligence System
from typing import Dict, Any

# Import Weekly Reporting Service
from weekly_reporting_service import WeeklyReportingService, test_weekly_report

# Import Partner Service
from partner_service import PartnerService, partner_service

# Advanced AI Intelligence System
class AIIntelligencePhases:
    """3-Phase AI Intelligence System for Construction Industry"""
    
    @staticmethod
    def get_enhanced_prompts() -> Dict[str, str]:
        """Phase 1: Enhanced Prompting - Master Instruction System for Construction Industry"""
        return {
            # MASTER SYSTEM PROMPT - FOUNDATION
            "general": """
You are ONESource AI, the definitive construction compliance advisor for AU/NZ markets.

ENHANCED SECTION FRAMEWORK - SELECTIVE USE ONLY:

ALWAYS INCLUDE (Core sections for every response):
🔧 **Technical Answer** - Comprehensive technical guidance with specific code references
🧐 **Mentoring Insight** - Professional development context and strategic guidance  
📋 **Next Steps** - Prioritized implementation roadmap

INCLUDE ONLY WHEN APPLICABLE:
📊 **Code Requirements** → only if multiple codes/standards apply
✅ **Compliance Verification** → only if verification process is complex
🔄 **Alternative Solutions** → only if suggesting a performance solution
🏛️ **Authority Requirements** → only if authority involvement is needed
📄 **Documentation Needed** → only if specific documentation is required
⚙️ **Workflow Recommendations** → only if workflow coordination is critical
🏗️ **Partner Intel** → only when verified community knowledge is available
🔍 **Partner Unverified** → only when unverified community insights exist
❓ **Clarifying Questions** → only if location/jurisdiction unknown after 2 related queries

NCC-FIRST HIERARCHY & COMPLIANCE RULES:
1. ALWAYS prioritize National Construction Code (NCC) requirements first
2. Reference AS/NZS standards as supporting documentation to NCC
3. Specify current code year (2025 NCC with amendments)
4. For conflicts: NCC supersedes Australian Standards
5. State/territory variations: Always note when applicable
6. Performance vs Deemed-to-Satisfy: Explain both pathways clearly

RESPONSE QUALITY STANDARDS:
- Professional, authoritative tone with clean formatting
- Specific, actionable recommendations
- Use tables ONLY when multiple options need comparison
- Clear implementation pathways
- Minimal section headers for clean presentation

CRITICAL: Always use 🧐 (professor with monocle) for Mentoring Insight sections.
CRITICAL: Only include sections that add value to the specific question asked.
""",

            # SPECIALIZED DISCIPLINES
            "hydraulic_engineering": """
You are a senior hydraulic engineer specializing in AU/NZ plumbing and drainage systems.

ENHANCED SECTION FRAMEWORK - MANDATORY STRUCTURE:
🔧 **Technical Answer** - Detailed hydraulic analysis with calculations
🧐 **Mentoring Insight** - Engineering best practices and professional development
📋 **Next Steps** - Systematic implementation sequence
📊 **Code Requirements** - NCC Section F and AS/NZS 3500 series compliance
✅ **Compliance Verification** - Testing, commissioning, and certification requirements
🏛️ **Authority Requirements** - Water authority approvals and council requirements
📄 **Documentation Needed** - Professional certifications and compliance documentation

NCC-FIRST HIERARCHY:
1. NCC Volume 1/2 Section F (Water services and facilities) - PRIMARY
2. AS/NZS 3500.1 (Water services) - Supporting standard
3. AS/NZS 3500.2 (Sanitary plumbing and drainage) - Supporting standard  
4. AS/NZS 3500.3 (Stormwater drainage) - Supporting standard
5. AS 2419 (Fire hydrant installations) - Where applicable

PROFESSIONAL SPECIALIZATION:
- Hydraulic calculations and system sizing
- Pipe sizing tables and flow rate analysis
- Pressure loss calculations and pump selection
- Backflow prevention and cross-connection control
- Water efficiency and WELS compliance
- Stormwater management and detention design

RESPONSE APPROACH:
- Include specific pipe sizing recommendations
- Provide hydraulic calculation methods
- Reference licensed plumber requirements
- Detail testing and commissioning procedures
- Address water authority approval processes
""",

            "structural_engineering": """
You are a senior structural engineer specializing in AU/NZ structural design and construction.

ENHANCED SECTION FRAMEWORK - MANDATORY STRUCTURE:
🔧 **Technical Answer** - Detailed structural analysis with load calculations
🧐 **Mentoring Insight** - Engineering judgment and professional development guidance
📋 **Next Steps** - Design and construction sequence
📊 **Code Requirements** - NCC structural provisions and AS/NZS structural standards
✅ **Compliance Verification** - Design review and certification requirements
🏛️ **Authority Requirements** - Building consent and structural certification
📄 **Documentation Needed** - Structural drawings, calculations, and certificates

NCC-FIRST HIERARCHY:
1. NCC Volume 1/2 Part B (Structure) - PRIMARY
2. AS/NZS 1170 series (Structural design actions) - Supporting standards
3. AS 3600 (Concrete structures) - Material-specific standard
4. AS 4100 (Steel structures) - Material-specific standard
5. AS 1720 (Timber structures) - Material-specific standard

PROFESSIONAL SPECIALIZATION:
- Wind load calculations per AS/NZS 1170.2
- Seismic design per AS/NZS 1170.4
- Concrete design per AS 3600
- Steel connection design per AS 4100
- Foundation design and geotechnical considerations
- Construction methodology and sequencing

RESPONSE APPROACH:
- Include specific load calculations
- Provide material design parameters
- Reference structural engineer certification requirements
- Detail construction tolerances and quality requirements
- Address peer review recommendations for complex projects
""",

            "fire_engineering": """
You are a fire safety engineer specializing in AU/NZ fire engineering and performance-based solutions.

ENHANCED SECTION FRAMEWORK - MANDATORY STRUCTURE:
🔧 **Technical Answer** - Fire engineering analysis with performance criteria
🧐 **Mentoring Insight** - Fire safety strategy and professional development
📋 **Next Steps** - Implementation and approval pathway
📊 **Code Requirements** - NCC fire safety provisions and AS fire standards
✅ **Compliance Verification** - Fire engineering report and authority approval
🏛️ **Authority Requirements** - Fire authority consultation and building certifier approval
📄 **Documentation Needed** - Fire engineering report, FRL certificates, maintenance schedules

NCC-FIRST HIERARCHY:
1. NCC Volume 1 Part C/E (Fire resistance and safety) - PRIMARY
2. NCC Volume 2 Part 3.7 (Fire safety) - Residential buildings
3. AS 1851 (Fire protection system maintenance) - Supporting standard
4. AS 2118 (Fire sprinkler systems) - Supporting standard
5. AS 1670 (Fire detection and alarm systems) - Supporting standard

PROFESSIONAL SPECIALIZATION:
- Performance-based fire safety design
- Fire engineering analysis and modeling
- Egress analysis and travel time calculations
- Fire resistance level (FRL) requirements
- Smoke management system design
- Fire brigade access and facilities

RESPONSE APPROACH:
- Include fire safety objectives and performance criteria
- Provide fire modeling parameters and assumptions
- Reference fire engineer certification requirements
- Detail testing and commissioning procedures
- Address fire authority consultation requirements
""",

            "building_surveying": """
You are a registered building surveyor specializing in AU/NZ building compliance and certification.

ENHANCED SECTION FRAMEWORK - MANDATORY STRUCTURE:
🔧 **Technical Answer** - Compliance pathway and regulatory requirements
🧐 **Mentoring Insight** - Regulatory strategy and professional development
📋 **Next Steps** - Building consent and certification process
📊 **Code Requirements** - NCC compliance requirements and assessment methods
✅ **Compliance Verification** - Inspection schedules and certification requirements
🏛️ **Authority Requirements** - Council/private certifier requirements and approvals
📄 **Documentation Needed** - Building consent documentation, compliance schedules, certificates

NCC-FIRST HIERARCHY:
1. National Construction Code (NCC) 2025 - PRIMARY reference
2. State/territory building regulations - Jurisdictional requirements
3. Australian Standards - Supporting technical standards
4. BCA Guides and Handbooks - Interpretation guidance
5. ABCB protocols and determinations - Official interpretations

PROFESSIONAL SPECIALIZATION:
- Building classification and use determinations
- NCC compliance assessment methodologies
- Alternative solution evaluation and approval
- Building consent application processes
- Construction inspection and certification
- Occupancy certificate and final approval procedures

RESPONSE APPROACH:
- Specify exact NCC clause references
- Explain compliance pathway options
- Detail building surveyor role and responsibilities
- Address inspection and certification schedules
- Reference professional indemnity and registration requirements
""",
        }
    
    @staticmethod
    def get_sector_contexts() -> Dict[str, str]:
        """Industry sector-specific contexts for 17 sectors"""
        return {
            "commercial": """
            SECTOR CONTEXT: Commercial buildings (offices, retail, mixed-use developments)
            KEY CONSIDERATIONS: Open plan flexibility, high occupancy loads, 24/7 operations, tenant requirements, base building vs tenant fit-out, accessibility compliance, energy efficiency ratings
            TYPICAL CHALLENGES: Services distribution in open plan, acoustic privacy, after-hours security, parking ratios, waste management
            """,
            
            "industrial": """
            SECTOR CONTEXT: Industrial facilities (warehouses, manufacturing plants, distribution hubs)
            KEY CONSIDERATIONS: Heavy loading, high bay construction, specialised services, process requirements, heavy vehicle access, dangerous goods storage
            TYPICAL CHALLENGES: Structural loads, fire protection for high rack storage, process ventilation, crane loads, truck maneuvering areas
            """,
            
            "mining_resources": """
            SECTOR CONTEXT: Mining infrastructure (processing plants, camps, remote facilities)
            KEY CONSIDERATIONS: Remote locations, harsh environments, temporary vs permanent facilities, worker accommodation, process plant requirements
            TYPICAL CHALLENGES: Extreme weather design, logistics constraints, power generation, water supply, waste disposal
            """,
            
            "health": """
            SECTOR CONTEXT: Healthcare facilities (hospitals, aged care, medical centres, specialist clinics)
            KEY CONSIDERATIONS: Infection control, life safety systems, medical gases, specialised equipment, patient accessibility, 24/7 operations
            TYPICAL CHALLENGES: Air pressure relationships, electrical reliability, water quality, medical equipment integration, patient flow
            """,
            
            "data_centres": """
            SECTOR CONTEXT: Data processing facilities (hyperscale, enterprise, colocation)
            KEY CONSIDERATIONS: Power density, cooling requirements, redundancy, security, electromagnetic compatibility, future expansion
            TYPICAL CHALLENGES: Power distribution, cooling efficiency, fire suppression (water-free), raised floor systems, cable management
            """,
            
            "residential": """
            SECTOR CONTEXT: Residential buildings (apartments, townhouses, detached housing)
            KEY CONSIDERATIONS: Acoustic privacy, accessibility, energy efficiency, natural ventilation, fire safety, building defects legislation
            TYPICAL CHALLENGES: Acoustic separation, waterproofing, thermal performance, fire egress, common area maintenance
            """,
            
            "hotels_hospitality": """
            SECTOR CONTEXT: Hospitality facilities (hotels, resorts, serviced apartments)
            KEY CONSIDERATIONS: Guest comfort, noise control, 24/7 operations, food service facilities, recreational amenities, security
            TYPICAL CHALLENGES: Acoustic privacy between rooms, hot water systems, kitchen ventilation, swimming pool areas, fire safety
            """,
            
            "education": """
            SECTOR CONTEXT: Educational facilities (schools, universities, research facilities)
            KEY CONSIDERATIONS: Large occupancy numbers, varied age groups, specialised teaching spaces, accessibility, security, future flexibility
            TYPICAL CHALLENGES: Assembly occupancy, laboratory requirements, sports facility integration, acoustic control, technology infrastructure
            """,
            
            "transport_infrastructure": """
            SECTOR CONTEXT: Transport facilities (airports, rail, ports, roads, bridges)
            KEY CONSIDERATIONS: Public access, high traffic volumes, security requirements, weather exposure, maintenance access, future expansion
            TYPICAL CHALLENGES: Structural loads, environmental exposure, public safety, operational continuity, heritage considerations
            """,
            
            "energy_utilities": """
            SECTOR CONTEXT: Energy facilities (power generation, renewable energy, water/wastewater plants)
            KEY CONSIDERATIONS: Process requirements, environmental impact, safety systems, regulatory compliance, operational reliability
            TYPICAL CHALLENGES: Hazardous area design, environmental compliance, process safety, maintenance access, emergency systems
            """,
            
            "government_civic": """
            SECTOR CONTEXT: Government facilities (courthouses, civic centres, defence facilities)
            KEY CONSIDERATIONS: Public access, security requirements, ceremonial spaces, accessibility compliance, heritage considerations
            TYPICAL CHALLENGES: Security vs accessibility balance, public gathering areas, technology integration, long-term durability
            """,
            
            "retail_entertainment": """
            SECTOR CONTEXT: Public facilities (shopping centres, cinemas, stadiums, cultural venues)
            KEY CONSIDERATIONS: High occupancy, emergency egress, accessibility, atmospheric control, entertainment systems, food service
            TYPICAL CHALLENGES: Large span structures, acoustic control, crowd management, fire safety in assembly spaces, HVAC for variable loads
            """,
            
            "agriculture_food": """
            SECTOR CONTEXT: Food processing facilities (abattoirs, cold storage, food manufacturing plants)
            KEY CONSIDERATIONS: Food safety requirements, temperature control, hygiene standards, process equipment, waste management
            TYPICAL CHALLENGES: Temperature and humidity control, drainage systems, equipment loads, cleaning requirements, pest exclusion
            """,
            
            "sports_recreation": """
            SECTOR CONTEXT: Recreation facilities (aquatic centres, gyms, sports complexes)
            KEY CONSIDERATIONS: Specialised environments, high humidity areas, equipment loads, spectator areas, accessibility
            TYPICAL CHALLENGES: Pool hall environments, equipment foundations, acoustic control, ventilation for high occupancy
            """,
            
            "mixed_use": """
            SECTOR CONTEXT: Mixed-use developments (integrated residential/commercial/retail hubs)
            KEY CONSIDERATIONS: Multiple occupancy types, shared services, acoustic separation, fire separation, parking requirements
            TYPICAL CHALLENGES: Services distribution, fire compartmentation, acoustic separation, waste management, access control
            """,
            
            "specialist_facilities": """
            SECTOR CONTEXT: Specialised facilities (laboratories, cleanrooms, high-security facilities)
            KEY CONSIDERATIONS: Controlled environments, specialised services, security requirements, contamination control, precise environmental control
            TYPICAL CHALLENGES: Air quality control, vibration control, electromagnetic compatibility, waste disposal, decontamination systems
            """,
            
            "others": """
            SECTOR CONTEXT: Other specialised or unique facility types not covered by standard categories
            KEY CONSIDERATIONS: Project-specific requirements, unique challenges, specialised regulatory requirements, innovative solutions
            TYPICAL CHALLENGES: Non-standard design criteria, limited precedents, specialised approval processes, unique risk factors
            """
        }
    
    @staticmethod 
    def detect_workflow_stage(question: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
        """Phase 2: Workflow Intelligence - Detect project stage and suggest next steps"""
        
        question_lower = question.lower()
        
        # Project stage detection
        if any(word in question_lower for word in ["planning", "concept", "feasibility", "initial"]):
            stage = "concept_planning"
        elif any(word in question_lower for word in ["design", "drawings", "architect", "plans"]):
            stage = "design_development"
        elif any(word in question_lower for word in ["approval", "consent", "permit", "certifier"]):
            stage = "regulatory_approval"  
        elif any(word in question_lower for word in ["tender", "contractor", "quote", "pricing"]):
            stage = "procurement"
        elif any(word in question_lower for word in ["construction", "building", "site", "concrete"]):
            stage = "construction"
        elif any(word in question_lower for word in ["inspection", "completion", "handover", "defects"]):
            stage = "completion"
        else:
            stage = "general_inquiry"
        
        # Workflow recommendations based on stage
        workflows = {
            "concept_planning": {
                "current_stage": "Concept & Planning",
                "typical_next_steps": [
                    "Engage architect/designer for preliminary concepts",
                    "Conduct site analysis and surveys",
                    "Preliminary budget estimation",
                    "Council pre-application advice",
                    "Geotechnical investigation if required"
                ],
                "key_consultants": ["Architect", "Town Planner", "Surveyor"],
                "critical_considerations": ["Zoning compliance", "Site constraints", "Budget parameters"]
            },
            "design_development": {
                "current_stage": "Design Development", 
                "typical_next_steps": [
                    "Detailed architectural drawings",
                    "Structural engineering design",
                    "Services engineering (mechanical, electrical, hydraulic)",
                    "Energy efficiency modeling",
                    "Accessibility compliance review"
                ],
                "key_consultants": ["Structural Engineer", "Services Engineer", "Energy Assessor"],
                "critical_considerations": ["NCC compliance", "Structural adequacy", "Energy efficiency"]
            },
            "regulatory_approval": {
                "current_stage": "Regulatory Approval",
                "typical_next_steps": [
                    "Building consent application preparation",
                    "Engineering calculations and certificates", 
                    "Fire safety report if required",
                    "Accessibility compliance statement",
                    "Council/certifier submission"
                ],
                "key_consultants": ["Building Certifier", "Fire Engineer", "Access Consultant"],
                "critical_considerations": ["Complete documentation", "Professional certifications", "Authority requirements"]
            },
            "procurement": {
                "current_stage": "Procurement & Tendering",
                "typical_next_steps": [
                    "Tender documentation preparation",
                    "Contractor selection and vetting",
                    "Contract negotiation and execution",
                    "Insurance and bonding arrangements",
                    "Construction program development"
                ],
                "key_consultants": ["Quantity Surveyor", "Contract Administrator", "Project Manager"],  
                "critical_considerations": ["Contract terms", "Insurance adequacy", "Quality assurance"]
            },
            "construction": {
                "current_stage": "Construction Phase",
                "typical_next_steps": [
                    "Regular site inspections and quality control",
                    "Progress payments and variation management", 
                    "Mandatory inspections scheduling",
                    "Material testing and compliance verification",
                    "Coordination of trades and services"
                ],
                "key_consultants": ["Site Supervisor", "Quality Assurance", "Testing Services"],
                "critical_considerations": ["Safety compliance", "Quality control", "Program adherence"]
            },
            "completion": {
                "current_stage": "Completion & Handover",
                "typical_next_steps": [
                    "Final inspections and compliance verification",
                    "Defects identification and rectification",
                    "Completion certificates and warranties",
                    "Operation and maintenance manual handover",
                    "Final account settlement"
                ],
                "key_consultants": ["Building Inspector", "Maintenance Contractor", "Warranty Provider"],
                "critical_considerations": ["Defects liability", "Warranty coverage", "Maintenance requirements"]
            },
            "general_inquiry": {
                "current_stage": "Information Gathering",
                "typical_next_steps": [
                    "Define project scope and objectives",
                    "Identify key stakeholders and consultants",
                    "Establish preliminary timeline and budget",
                    "Research applicable standards and regulations"
                ],
                "key_consultants": ["Project Advisor", "Relevant Specialist"],
                "critical_considerations": ["Scope definition", "Resource planning", "Regulatory research"]
            }
        }
        
        return workflows.get(stage, workflows["general_inquiry"])
    
    @staticmethod
    def get_specialized_context(discipline: str, question: str) -> Dict[str, Any]:
        """Phase 3: Specialized Training - Discipline-specific knowledge enhancement"""
        
        specialized_knowledge = {
            "structural": {
                "key_standards": [
                    "AS 1170.0 - Structural design actions - General principles",
                    "AS 1170.1 - Permanent, imposed and other actions", 
                    "AS 1170.2 - Wind actions",
                    "AS 1170.4 - Earthquake actions",
                    "AS 3600 - Concrete structures",
                    "AS 4100 - Steel structures"
                ],
                "common_calculations": [
                    "Wind load calculations per AS 1170.2",
                    "Seismic design per AS 1170.4", 
                    "Concrete design per AS 3600",
                    "Steel connection design per AS 4100"
                ],
                "professional_requirements": [
                    "Structural engineer certification required",
                    "Professional indemnity insurance essential", 
                    "Regular CPD maintenance required",
                    "Peer review recommended for complex projects"
                ]
            },
            
            "fire_safety": {
                "key_standards": [
                    "AS 1530 - Methods for fire tests on building materials",
                    "AS 1851 - Maintenance of fire protection systems",
                    "AS 2118 - Automatic fire sprinkler systems", 
                    "AS 3786 - Smoke alarms using scattered light",
                    "AS 4072 - Components for fire detection systems"
                ],
                "design_considerations": [
                    "Building classification and fire safety objectives",
                    "Egress analysis and travel distances",
                    "Fire resistance levels (FRL) requirements",
                    "Smoke hazard management systems"
                ],
                "compliance_verification": [
                    "Fire engineering report required for performance solutions",
                    "Fire authority consultation for complex buildings",
                    "Third-party certification for critical systems"
                ]
            },
            
            "mechanical": {
                "key_standards": [
                    "AS 1668 - The use of mechanical ventilation",
                    "AS 3700 - Masonry structures", 
                    "AS 5601 - Gas installations",
                    "AS/NZS 3000 - Electrical installations"
                ],
                "system_design": [
                    "HVAC load calculations and equipment sizing",
                    "Ventilation rates per AS 1668",
                    "Energy efficiency per NCC Section J",
                    "Refrigerant selection and environmental impact"
                ],
                "installation_requirements": [
                    "Licensed tradesperson installation mandatory",
                    "Pressure testing and commissioning required",
                    "Operation and maintenance manual provision"
                ]
            },
            
            "hydraulic": {
                "key_standards": [
                    "AS/NZS 3500 - Plumbing and drainage",
                    "AS 2419 - Fire hydrant installations",
                    "AS 3500.1 - Water services",
                    "AS 3500.2 - Sanitary plumbing and drainage"  
                ],
                "design_principles": [
                    "Water supply sizing and pressure requirements",
                    "Drainage design and pipe sizing",
                    "Stormwater management and detention",
                    "Water efficiency and WELS compliance"
                ],
                "regulatory_aspects": [
                    "Licensed plumber installation required",
                    "Water authority approvals for connections",
                    "Backflow prevention device mandatory"
                ]
            }
        }
        
        # Detect discipline from question content - ENHANCED DETECTION
        question_lower = question.lower()
        detected_discipline = "building_codes"  # Default
        
        # MECHANICAL & BUILDING SERVICES
        if any(word in question_lower for word in ["hydraulic", "plumbing", "water supply", "drainage", "sewer", "pipes"]):
            detected_discipline = "hydraulic_engineering"
        elif any(word in question_lower for word in ["hvac", "mechanical", "ventilation", "heating", "cooling", "air conditioning"]):
            detected_discipline = "mechanical_services"
        elif any(word in question_lower for word in ["sprinkler", "fire pump", "suppression", "wet system", "dry system"]):
            detected_discipline = "fire_protection_wet_dry"
        elif any(word in question_lower for word in ["fire engineering", "performance based", "fire model", "egress"]):
            detected_discipline = "fire_engineering"
        elif any(word in question_lower for word in ["electrical", "power", "lighting", "switchboard", "cable"]):
            detected_discipline = "electrical_engineering"
        elif any(word in question_lower for word in ["communications", "ict", "data", "telecommunications", "network"]):
            detected_discipline = "communications_ict"
        elif any(word in question_lower for word in ["security", "access control", "cctv", "alarm"]):
            detected_discipline = "security_systems"
        elif any(word in question_lower for word in ["lift", "elevator", "escalator", "vertical transport"]):
            detected_discipline = "vertical_transportation"
        elif any(word in question_lower for word in ["bms", "building automation", "control system", "smart building"]):
            detected_discipline = "building_automation"
        elif any(word in question_lower for word in ["acoustic", "noise", "sound", "vibration"]):
            detected_discipline = "acoustic_engineering"
        elif any(word in question_lower for word in ["lighting design", "illumination", "daylight"]):
            detected_discipline = "lighting_design"
            
        # STRUCTURAL & CIVIL
        elif any(word in question_lower for word in ["structural", "beam", "column", "foundation", "load", "concrete", "steel"]):
            detected_discipline = "structural_engineering"
        elif any(word in question_lower for word in ["civil", "site development", "earthworks", "pavement", "infrastructure"]):
            detected_discipline = "civil_engineering"
        elif any(word in question_lower for word in ["geotechnical", "soil", "foundation", "pile", "bearing capacity"]):
            detected_discipline = "geotechnical_engineering"
        elif any(word in question_lower for word in ["seismic", "earthquake", "dynamic", "ductility"]):
            detected_discipline = "seismic_engineering"
            
        # ARCHITECTURE & DESIGN
        elif any(word in question_lower for word in ["architecture", "design", "space planning", "layout"]):
            detected_discipline = "architecture"
        elif any(word in question_lower for word in ["interior design", "fitout", "furniture", "finishes"]):
            detected_discipline = "interior_design"
        elif any(word in question_lower for word in ["landscape", "outdoor", "garden", "plaza"]):
            detected_discipline = "landscape_architecture"
        elif any(word in question_lower for word in ["urban planning", "town planning", "zoning"]):
            detected_discipline = "urban_design"
            
        # SUSTAINABILITY & ENVIRONMENTAL
        elif any(word in question_lower for word in ["sustainability", "green building", "environmental", "carbon"]):
            detected_discipline = "sustainability"
        elif any(word in question_lower for word in ["energy modelling", "energy efficiency", "thermal performance"]):
            detected_discipline = "energy_modelling"
            
        # BUILDING ENVELOPE
        elif any(word in question_lower for word in ["facade", "curtain wall", "glazing", "window"]):
            detected_discipline = "facade_engineering"
        elif any(word in question_lower for word in ["waterproofing", "membrane", "sealant", "water ingress"]):
            detected_discipline = "waterproofing_design"
        elif any(word in question_lower for word in ["roofing", "roof", "membrane", "tiles"]):
            detected_discipline = "roofing_systems"
        elif any(word in question_lower for word in ["cladding", "external wall", "panels"]):
            detected_discipline = "cladding_systems"
        elif any(word in question_lower for word in ["green star", "nabers", "well", "certification"]):
            detected_discipline = "green_building_certification"
            
        # PROJECT MANAGEMENT & COMMERCIAL
        elif any(word in question_lower for word in ["project management", "program", "schedule", "coordination"]):
            detected_discipline = "project_management"
        elif any(word in question_lower for word in ["cost", "quantity surveying", "estimate", "budget"]):
            detected_discipline = "cost_planning"
        elif any(word in question_lower for word in ["contract", "legal", "procurement", "tender"]):
            detected_discipline = "contract_administration"
        elif any(word in question_lower for word in ["design management", "coordination", "BIM"]):
            detected_discipline = "design_management"
        elif any(word in question_lower for word in ["risk", "safety", "hazard"]):
            detected_discipline = "risk_management"
            
        # SPECIALIST DISCIPLINES
        elif any(word in question_lower for word in ["heritage", "conservation", "restoration"]):
            detected_discipline = "heritage_conservation"
        elif any(word in question_lower for word in ["modular", "prefab", "offsite"]):
            detected_discipline = "modular_prefabrication"
        elif any(word in question_lower for word in ["industrial", "process", "manufacturing", "plant"]):
            detected_discipline = "industrial_process"
        elif any(word in question_lower for word in ["asbestos", "lead", "hazardous materials"]):
            detected_discipline = "hazardous_materials"
        elif any(word in question_lower for word in ["environmental engineering", "contamination", "remediation"]):
            detected_discipline = "environmental_engineering"
        elif any(word in question_lower for word in ["traffic", "transport", "parking", "roads"]):
            detected_discipline = "traffic_transport"
        elif any(word in question_lower for word in ["health planning", "medical", "hospital", "healthcare"]):
            detected_discipline = "health_planning"
        elif any(word in question_lower for word in ["wayfinding", "signage", "navigation"]):
            detected_discipline = "wayfinding_signage"
        elif any(word in question_lower for word in ["waste", "recycling", "garbage"]):
            detected_discipline = "waste_management"
        
        # Detect sector from question content
        detected_sector = "others"  # Default
        
        if any(word in question_lower for word in ["office", "commercial", "retail", "mixed use"]):
            detected_sector = "commercial"
        elif any(word in question_lower for word in ["warehouse", "industrial", "manufacturing", "factory"]):
            detected_sector = "industrial"
        elif any(word in question_lower for word in ["mining", "mine", "resources", "camp"]):
            detected_sector = "mining_resources"  
        elif any(word in question_lower for word in ["hospital", "medical", "health", "aged care", "clinic"]):
            detected_sector = "health"
        elif any(word in question_lower for word in ["data centre", "server", "data center"]):
            detected_sector = "data_centres"
        elif any(word in question_lower for word in ["residential", "apartment", "house", "townhouse"]):
            detected_sector = "residential"
        elif any(word in question_lower for word in ["hotel", "hospitality", "resort", "accommodation"]):
            detected_sector = "hotels_hospitality"
        elif any(word in question_lower for word in ["school", "university", "education", "classroom"]):
            detected_sector = "education"
        elif any(word in question_lower for word in ["airport", "rail", "transport", "infrastructure", "bridge"]):
            detected_sector = "transport_infrastructure"
        elif any(word in question_lower for word in ["power plant", "utility", "energy", "water treatment"]):
            detected_sector = "energy_utilities"
        elif any(word in question_lower for word in ["government", "civic", "courthouse", "defence"]):
            detected_sector = "government_civic"
        elif any(word in question_lower for word in ["shopping", "cinema", "stadium", "entertainment"]):
            detected_sector = "retail_entertainment"
        elif any(word in question_lower for word in ["food processing", "abattoir", "cold storage"]):
            detected_sector = "agriculture_food"
        elif any(word in question_lower for word in ["sports", "recreation", "gym", "aquatic"]):
            detected_sector = "sports_recreation"
        elif any(word in question_lower for word in ["laboratory", "cleanroom", "security facility"]):
            detected_sector = "specialist_facilities"
            
        return {
            "detected_discipline": detected_discipline,
            "specialized_knowledge": specialized_knowledge.get(detected_discipline, {}),
            "cross_discipline_considerations": [
                "Coordination with other engineering disciplines required",
                "Integrated design approach recommended", 
                "Professional liability and insurance considerations",
                "Quality assurance and peer review processes"
            ]
        }

# Import our services
from firebase_service import firebase_service
from payment_service import PaymentService

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Import AI service AFTER loading environment variables
from ai_service import construction_ai

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Initialize OpenAI client
openai_client = AsyncOpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

# Initialize payment service
payment_service = PaymentService(db)

# Create the main app without a prefix
app = FastAPI(title="ONESource-ai API", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()

# Pydantic Models
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

class UserOnboardingData(BaseModel):
    name: str
    profession: str
    sector: str
    use_case: str
    marketing_consent: bool = False

class ChatQuestion(BaseModel):
    question: str
    session_id: Optional[str] = None

class ChatFeedback(BaseModel):
    message_id: str
    feedback_type: str  # 'positive' or 'negative'
    comment: Optional[str] = None

class KnowledgeContribution(BaseModel):
    message_id: str
    contribution: str
    opt_in_credit: bool = True

class ChatHistoryRequest(BaseModel):
    limit: int = 50

# Knowledge Management Models
class DocumentUpload(BaseModel):
    file_name: str
    file_type: str
    content_type: str
    tags: List[str] = []
    is_supplier_content: bool = False
    supplier_info: Optional[Dict[str, Any]] = None

class SupplierInfo(BaseModel):
    company_name: str
    abn: Optional[str] = None
    contact_email: str
    logo_url: Optional[str] = None
    product_tags: List[str] = []
    terms_agreed: bool = False

class MentorNote(BaseModel):
    title: str
    content: str
    tags: List[str] = []
    category: Optional[str] = None
    attachment_url: Optional[str] = None

class KnowledgeSearch(BaseModel):
    query: str
    limit: int = 10
    include_supplier_only: bool = False

class PaymentRequest(BaseModel):
    package_id: str
    origin_url: str

class VoucherRequest(BaseModel):
    voucher_code: str

class VoucherCreate(BaseModel):
    voucher_code: str
    plan_type: str  # 'pro', 'consultant', 'day_pass'
    duration_days: int = 30
    max_uses: int = 1
    description: Optional[str] = None

class UserPreferences(BaseModel):
    industries: List[str] = []
    role: str = ""
    experience_level: str = ""
    response_style: str = "balanced"
    ai_focus_areas: List[str] = []
    custom_instructions: str = ""

class PartnerRegistration(BaseModel):
    company_name: str
    abn: str
    primary_contact_name: str
    primary_email: str
    backup_email: str
    agreed_to_terms: bool
    description: Optional[str] = None

# Authentication dependency
async def get_current_user(request: Request, credentials = Depends(security)) -> Dict[str, Any]:
    """Verify Firebase token and return user info"""
    try:
        token = credentials.credentials
        user_info = await firebase_service.verify_token(token)
        if not user_info:
            raise HTTPException(status_code=401, detail="Invalid authentication token")
        return user_info
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Authentication failed: {str(e)}")

# Optional authentication dependency
async def get_current_user_optional(request: Request) -> Optional[Dict[str, Any]]:
    """Get user info if authenticated, otherwise None"""
    try:
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            return None
        
        token = auth_header.split(" ")[1]
        user_info = await firebase_service.verify_token(token)
        return user_info
    except:
        return None

# Basic routes
@api_router.get("/")
async def root():
    return {"message": "ONESource-ai API is running", "version": "1.0.0"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.dict()
    status_obj = StatusCheck(**status_dict)
    _ = await db.status_checks.insert_one(status_obj.dict())
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await db.status_checks.find().to_list(1000)
    return [StatusCheck(**status_check) for status_check in status_checks]

# User Management Routes
@api_router.post("/user/onboarding")
async def complete_onboarding(
    data: UserOnboardingData,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Complete user onboarding after authentication"""
    try:
        uid = current_user["uid"]
        
        # Prepare user profile data
        profile_data = {
            "name": data.name,
            "profession": data.profession,
            "sector": data.sector,
            "use_case": data.use_case,
            "marketing_consent": data.marketing_consent,
            "onboarding_completed": True,
            "subscription_tier": "starter",
            "trial_questions_used": 0,
            "subscription_active": False,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        # Save to Firebase
        success = await firebase_service.create_user_profile(uid, profile_data)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to save user profile")
        
        return {"message": "Onboarding completed successfully", "profile": profile_data}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error completing onboarding: {str(e)}")

@api_router.get("/user/profile")
async def get_user_profile(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get current user's profile"""
    try:
        uid = current_user["uid"]
        profile = await firebase_service.get_user_profile(uid)
        
        if not profile:
            # New user - return basic info from Firebase auth
            return {
                "uid": uid,
                "email": current_user.get("email"),
                "name": current_user.get("name"),
                "onboarding_completed": False,
                "subscription_tier": "starter",
                "trial_questions_used": 0
            }
        
        return profile
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting profile: {str(e)}")

@api_router.get("/user/subscription")
async def get_subscription_status(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Check user's subscription status and usage"""
    try:
        uid = current_user["uid"]
        subscription = await firebase_service.check_user_subscription(uid)
        
        # Add tier field for frontend compatibility
        subscription["tier"] = subscription.get("subscription_tier", "starter")
        
        # CRITICAL FIX: Transform for frontend compatibility - Pro users should NOT be in trial mode
        subscription_tier = subscription.get("subscription_tier", "starter")
        subscription_active = subscription.get("subscription_active", False)
        
        if subscription_tier == "starter":
            # Only starter users can be in trial mode
            subscription["is_trial"] = not subscription_active
            if subscription["is_trial"]:
                subscription["trial_info"] = {
                    "questions_remaining": max(0, 3 - subscription.get("trial_questions_used", 0)),
                    "questions_used": subscription.get("trial_questions_used", 0)
                }
        else:
            # Pro, Pro-Plus, and other paid tiers are NEVER in trial mode
            subscription["is_trial"] = False
            # Remove trial_info for paid users
            if "trial_info" in subscription:
                del subscription["trial_info"]
        
        return subscription
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error checking subscription: {str(e)}")

# AI Chat Routes
@api_router.post("/chat/ask")
async def ask_question(
    chat_data: ChatQuestion,
    request: Request,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user_optional)
):
    """Ask a construction industry question - UNIFIED BACKEND"""
    try:
        # Import shared service
        from shared_chat_service import shared_chat_service
        
        # Validate question is construction-related
        if not await construction_ai.validate_construction_question(chat_data.question):
            raise HTTPException(
                status_code=400, 
                detail="Questions must be related to AU/NZ construction industry"
            )
        
        # Check trial limits for unauthenticated users
        if not current_user:
            # Anonymous user - check if within trial limit
            session_id = chat_data.session_id or str(uuid.uuid4())
            user_profile = None
            trial_warning = True
        else:
            # Authenticated user - check subscription status
            uid = current_user["uid"]
            subscription = await firebase_service.check_user_subscription(uid)
            trial_questions_used = subscription["trial_questions_used"]
            
            # Check if user can ask questions
            if (subscription["subscription_tier"] == "starter" and 
                not subscription["subscription_active"]):
                
                # Get today's date for daily limit checking
                today = datetime.utcnow().date()
                
                # Get or create today's usage tracking
                daily_usage_key = f"daily_questions_{today.strftime('%Y%m%d')}"
                daily_questions_used = subscription.get(daily_usage_key, 0)
                
                if daily_questions_used >= 3:
                    raise HTTPException(
                        status_code=403,
                        detail={
                            "type": "daily_limit_reached",
                            "message": "You've reached your daily limit of 3 questions. Upgrade to ask unlimited questions.",
                            "questions_used": daily_questions_used,
                            "daily_limit": 3
                        }
                    )
                
                trial_warning = True
                remaining_questions = 3 - (daily_questions_used + 1)
                session_id = chat_data.session_id or str(uuid.uuid4())
                user_profile = None
            else:
                trial_warning = False
                remaining_questions = None
                session_id = chat_data.session_id or str(uuid.uuid4())
                user_profile = None
        
        # **UNIFIED BACKEND CALL - NO LEGACY LOGIC**
        response_data = shared_chat_service.get_unified_chat_response(
            question=chat_data.question,
            session_id=session_id,
            user_context=None  # Regular endpoint doesn't use enhanced features
        )
        
        # Add trial and subscription information  
        if trial_warning and not current_user:
            response_data["subscription_status"] = {
                "is_trial": True,
                "subscription_tier": "starter",
                "subscription_active": False,
                "trial_info": {
                    "questions_remaining": 2,  # Decremented after this question
                    "questions_used": 1
                }
            }
        elif trial_warning and current_user:
            response_data["subscription_status"] = {
                "is_trial": True,
                "subscription_tier": "starter", 
                "subscription_active": False,
                "trial_info": {
                    "questions_remaining": remaining_questions,
                    "questions_used": 3 - remaining_questions
                }
            }
        elif current_user:
            # Get actual subscription status for authenticated users
            uid = current_user["uid"]
            subscription = await firebase_service.check_user_subscription(uid)
            response_data["subscription_status"] = subscription
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in ask_question: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

# Document Processing and AI Functions
async def extract_text_from_file(file_content: bytes, content_type: str, filename: str) -> str:
    """Extract text content from uploaded files using AI and specialized libraries"""
    try:
        # For PDFs, use PyPDF2 for better text extraction
        if content_type == 'application/pdf':
            try:
                import PyPDF2
                import io
                
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"Page {page_num + 1}:\n{page_text}")
                
                extracted_text = "\n\n".join(text_content)
                
                # If no text was extracted or very little, use AI vision as fallback
                if len(extracted_text.strip()) < 50:
                    return f"PDF text extraction yielded minimal content from {filename}. Consider using OCR for scanned documents."
                
                return extracted_text
                
            except Exception as pdf_error:
                return f"Error extracting PDF content from {filename}: {str(pdf_error)}"
        
        # For Word documents, use python-docx
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            try:
                from docx import Document
                import io
                
                doc_file = io.BytesIO(file_content)
                doc = Document(doc_file)
                
                # Extract text from paragraphs
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                extracted_text = "\n".join(text_content)
                
                if len(extracted_text.strip()) < 10:
                    return f"Word document content extraction yielded no readable text from {filename}"
                
                return extracted_text
                
            except Exception as word_error:
                return f"Error extracting Word document content from {filename}: {str(word_error)}"
        
        # For images, use OpenAI Vision API
        elif content_type.startswith('image/'):
            try:
                base64_content = base64.b64encode(file_content).decode('utf-8')
                
                response = await openai_client.chat.completions.create(
                    model="gpt-4-vision-preview",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all text content from this construction document image. Include technical specifications, dimensions, standards references, supplier information, and any visible text. If this appears to be a construction drawing or blueprint, describe the key elements and any text/labels visible."},
                                {"type": "image_url", "image_url": {"url": f"data:{content_type};base64,{base64_content}"}}
                            ]
                        }
                    ],
                    max_tokens=2000
                )
                return response.choices[0].message.content
                
            except Exception as vision_error:
                return f"Error extracting content from image {filename}: {str(vision_error)}"
        
        # For plain text files and other formats
        else:
            try:
                # Try UTF-8 first
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    # Fallback to latin-1
                    return file_content.decode('latin-1')
                except UnicodeDecodeError:
                    return f"Unable to decode text content from {filename}. File may be in an unsupported encoding."
            
    except Exception as e:
        return f"Error processing file {filename}: {str(e)}"

async def generate_embeddings(text: str) -> List[float]:
    """Generate vector embeddings for text using OpenAI"""
    try:
        # Check if we have a valid OpenAI API key
        api_key = os.environ.get('OPENAI_API_KEY', '')
        if not api_key or len(api_key) < 10:
            # Mock embeddings for testing - generate a simple hash-based embedding
            import hashlib
            text_hash = hashlib.md5(text.encode()).hexdigest()
            # Convert hex to float array (1536 dimensions like text-embedding-ada-002)
            mock_embedding = []
            for i in range(0, len(text_hash), 2):
                hex_pair = text_hash[i:i+2]
                mock_embedding.append(int(hex_pair, 16) / 255.0 - 0.5)  # Normalize to [-0.5, 0.5]
            
            # Pad or truncate to 1536 dimensions
            while len(mock_embedding) < 1536:
                mock_embedding.extend(mock_embedding[:min(len(mock_embedding), 1536 - len(mock_embedding))])
            mock_embedding = mock_embedding[:1536]
            
            return mock_embedding
        
        response = await openai_client.embeddings.create(
            model="text-embedding-ada-002",
            input=text[:8000]  # Limit to token constraints
        )
        return response.data[0].embedding
    except Exception as e:
        print(f"Error generating embeddings: {e}")
        # Return mock embedding on error
        return [0.1] * 1536

async def parse_document_metadata(text_content: str, filename: str, is_supplier: bool = False) -> Dict[str, Any]:
    """AI-powered extraction of document metadata and tags"""
    try:
        # Check if we have a valid OpenAI API key
        api_key = os.environ.get('OPENAI_API_KEY', '')
        if not api_key or len(api_key) < 10:
            # Mock metadata extraction for testing
            mock_tags = ["construction"]
            if "steel" in text_content.lower() or "beam" in text_content.lower():
                mock_tags.extend(["steel", "structural"])
            if "fire" in text_content.lower():
                mock_tags.append("fire-safety")
            if "hvac" in text_content.lower() or "ventilation" in text_content.lower():
                mock_tags.extend(["hvac", "mechanical"])
            if "AS/NZS" in text_content or "AS " in text_content:
                mock_tags.append("standards")
            
            return {
                "tags": mock_tags,
                "document_type": "specification" if is_supplier else "standard",
                "topics": [filename.replace("_", " ").replace(".txt", "")],
                "supplier_mentions": ["ACME Construction Materials"] if is_supplier else [],
                "categories": ["commercial" if "commercial" in text_content.lower() else "general"],
                "summary": f"Mock analysis of {filename}: Construction document with relevant technical content."
            }
        
        system_prompt = f"""
        Analyze this construction document and extract key metadata:
        
        1. Technical tags (building codes, standards, materials, etc.)
        2. Document type (specification, drawing, manual, etc.)
        3. Key topics and subjects
        4. Supplier information (if any company names mentioned)
        5. Construction categories (residential, commercial, industrial, etc.)
        
        Document filename: {filename}
        Is supplier content: {is_supplier}
        
        Return JSON format with: tags, document_type, topics, supplier_mentions, categories, summary
        """
        
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text_content[:4000]}  # Limit content size
            ],
            max_tokens=800
        )
        
        # Parse the JSON response
        import json
        try:
            metadata = json.loads(response.choices[0].message.content)
            return metadata
        except:
            # Fallback if JSON parsing fails
            return {
                "tags": ["construction", "document"],
                "document_type": "unknown",
                "topics": [filename],
                "supplier_mentions": [],
                "categories": ["general"],
                "summary": response.choices[0].message.content[:200]
            }
            
    except Exception as e:
        print(f"Error parsing document metadata: {e}")
        return {
            "tags": ["construction"],
            "document_type": "unknown", 
            "topics": [filename],
            "supplier_mentions": [],
            "categories": ["general"],
            "summary": "Document processing error"
        }

async def intelligent_knowledge_search(query: str, limit: int = 10) -> List[Dict[str, Any]]:
    """Search knowledge base using semantic similarity"""
    try:
        # Generate embedding for the query
        query_embedding = await generate_embeddings(query)
        if not query_embedding:
            return []
        
        # Search in knowledge vault
        cursor = db.knowledge_vault.find({})
        documents = await cursor.to_list(length=1000)  # Get all for similarity comparison
        
        scored_docs = []
        
        for doc in documents:
            if 'embedding' in doc and doc['embedding']:
                # Calculate cosine similarity
                doc_embedding = np.array(doc['embedding']).reshape(1, -1)
                query_emb = np.array(query_embedding).reshape(1, -1)
                similarity = cosine_similarity(query_emb, doc_embedding)[0][0]
                
                # Boost supplier content
                boost = 1.2 if doc.get('is_supplier_content', False) else 1.0
                final_score = similarity * boost
                
                # Clean up MongoDB ObjectId for JSON serialization
                doc_clean = dict(doc)
                if '_id' in doc_clean:
                    doc_clean['_id'] = str(doc_clean['_id'])
                
                scored_docs.append({
                    'document': doc_clean,
                    'similarity_score': final_score,
                    'is_supplier': doc.get('is_supplier_content', False)
                })
        
        # Sort by similarity score and return top results
        scored_docs.sort(key=lambda x: x['similarity_score'], reverse=True)
        return scored_docs[:limit]
        
    except Exception as e:
        print(f"Error in knowledge search: {e}")
        return []

# Knowledge Bank Routes - Two-Tier System
@api_router.post("/knowledge/upload-community")
async def upload_to_community_knowledge_bank(
    file: UploadFile = File(...),
    tags: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Upload document to Community Knowledge Bank (Partners and Admins only)"""
    try:
        uid = current_user["uid"]
        email = current_user.get("email")
        
        # Check if user is a partner or admin
        partner = await partner_service.get_partner_by_email(email) if email else None
        
        # For now, simple admin check - in production, implement proper role checking
        is_admin = email and "@onesource-ai.com" in email  # Simple admin check
        
        if not partner and not is_admin:
            raise HTTPException(
                status_code=403, 
                detail="Access denied. Only registered partners and administrators can upload to Community Knowledge Bank."
            )
        
        # Read file content
        file_content = await file.read()
        content_type = file.content_type or mimetypes.guess_type(file.filename)[0]
        
        # Generate file hash for deduplication
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Check for duplicates in community knowledge bank
        existing = await db.community_knowledge_bank.find_one({"file_hash": file_hash})
        if existing:
            raise HTTPException(status_code=400, detail="Document already exists in Community Knowledge Bank")
        
        # Extract text content using AI
        extracted_text = await extract_text_from_file(file_content, content_type, file.filename)
        
        # Generate AI metadata and tags
        metadata = await parse_document_metadata(extracted_text, file.filename, True)  # Always supplier content for community
        
        # Generate embeddings for semantic search
        embedding = await generate_embeddings(extracted_text)
        
        # Prepare document record for Community Knowledge Bank
        document_record = {
            "document_id": str(uuid.uuid4()),
            "user_id": uid,
            "uploader_email": email,
            "filename": file.filename,
            "original_size": len(file_content),
            "content_type": content_type,
            "file_hash": file_hash,
            "extracted_text": extracted_text,
            "embedding": embedding,
            "ai_metadata": metadata,
            "tags": tags.split(",") if tags else metadata.get("tags", []),
            "knowledge_bank_type": "community",  # Key distinction
            "partner_info": {
                "company_name": partner["company_name"],
                "partner_id": partner["partner_id"],
                "abn": partner["abn"]
            } if partner else None,
            "uploaded_by_admin": is_admin,
            "upload_timestamp": datetime.utcnow(),
            "status": "active",
            "view_count": 0,
            "reference_count": 0
        }
        
        # Store file reference
        try:
            file_base64 = base64.b64encode(file_content).decode('utf-8')
            storage_path = f"community_knowledge_bank/{document_record['document_id']}/{file.filename}"
            document_record["storage_path"] = storage_path
            document_record["file_data"] = file_base64[:1000]  # Store sample for testing
        except Exception as storage_error:
            print(f"Storage warning: {storage_error}")
            document_record["storage_path"] = f"local_storage/{document_record['document_id']}"
        
        # Save to Community Knowledge Bank collection
        await db.community_knowledge_bank.insert_one(document_record)
        
        # Update partner upload count
        if partner:
            await partner_service.increment_upload_count(partner["partner_id"])
            
            # Send upload receipt email
            document_info = {
                "filename": file.filename,
                "document_id": document_record["document_id"],
                "upload_date": document_record["upload_timestamp"].strftime("%d %B %Y at %H:%M AEDT"),
                "file_size": f"{len(file_content) / 1024:.1f} KB",
                "tags": document_record["tags"]
            }
            await partner_service.send_upload_receipt_email(partner, document_info)
        
        return {
            "message": "Document uploaded to Community Knowledge Bank successfully",
            "document_id": document_record["document_id"],
            "knowledge_bank": "community",
            "company_attribution": partner["company_name"] if partner else "ONESource-ai Admin",
            "extracted_summary": metadata.get("summary", ""),
            "detected_tags": metadata.get("tags", []),
            "email_receipt_sent": bool(partner)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading to Community Knowledge Bank: {str(e)}")

@api_router.post("/knowledge/upload-personal")
async def upload_to_personal_knowledge_bank(
    file: UploadFile = File(...),
    tags: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Upload document to Personal Knowledge Bank (Private to user)"""
    try:
        uid = current_user["uid"]
        
        # Read file content
        file_content = await file.read()
        content_type = file.content_type or mimetypes.guess_type(file.filename)[0]
        
        # Generate file hash for deduplication within user's personal bank
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Check for duplicates in user's personal knowledge bank
        existing = await db.personal_knowledge_bank.find_one({
            "user_id": uid,
            "file_hash": file_hash
        })
        if existing:
            raise HTTPException(status_code=400, detail="Document already exists in your Personal Knowledge Bank")
        
        # Extract text content using AI
        extracted_text = await extract_text_from_file(file_content, content_type, file.filename)
        
        # Generate AI metadata and tags
        metadata = await parse_document_metadata(extracted_text, file.filename, False)  # Not supplier content
        
        # Generate embeddings for semantic search
        embedding = await generate_embeddings(extracted_text)
        
        # Prepare document record for Personal Knowledge Bank
        document_record = {
            "document_id": str(uuid.uuid4()),
            "user_id": uid,
            "filename": file.filename,
            "original_size": len(file_content),
            "content_type": content_type,
            "file_hash": file_hash,
            "extracted_text": extracted_text,
            "embedding": embedding,
            "ai_metadata": metadata,
            "tags": tags.split(",") if tags else metadata.get("tags", []),
            "knowledge_bank_type": "personal",  # Key distinction
            "upload_timestamp": datetime.utcnow(),
            "status": "active",
            "view_count": 0,
            "reference_count": 0
        }
        
        # Store file reference
        try:
            file_base64 = base64.b64encode(file_content).decode('utf-8')
            storage_path = f"personal_knowledge_bank/{uid}/{document_record['document_id']}/{file.filename}"
            document_record["storage_path"] = storage_path
            document_record["file_data"] = file_base64[:1000]  # Store sample for testing
        except Exception as storage_error:
            print(f"Storage warning: {storage_error}")
            document_record["storage_path"] = f"local_storage/{document_record['document_id']}"
        
        # Save to Personal Knowledge Bank collection
        await db.personal_knowledge_bank.insert_one(document_record)
        
        return {
            "message": "Document uploaded to Personal Knowledge Bank successfully",
            "document_id": document_record["document_id"],
            "knowledge_bank": "personal",
            "privacy": "Private to your account only",
            "extracted_summary": metadata.get("summary", ""),
            "detected_tags": metadata.get("tags", [])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading to Personal Knowledge Bank: {str(e)}")

# Legacy endpoint - deprecated but kept for backward compatibility
@api_router.post("/knowledge/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    tags: str = Form(""),
    is_supplier_content: bool = Form(False),
    supplier_name: str = Form(""),
    supplier_abn: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """DEPRECATED: Legacy upload endpoint - use /upload-community or /upload-personal instead"""
    try:
        # For backward compatibility, route to personal knowledge bank
        return await upload_to_personal_knowledge_bank(file, tags, current_user)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading document: {str(e)}")

@api_router.post("/knowledge/mentor-note")
async def create_mentor_note(
    note_data: MentorNote,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Create a mentor note (enhanced wiki contribution)"""
    try:
        uid = current_user["uid"]
        
        # Generate embeddings for the note content
        full_text = f"{note_data.title} {note_data.content}"
        embedding = await generate_embeddings(full_text)
        
        # AI-powered categorization and tagging
        metadata = await parse_document_metadata(note_data.content, note_data.title, False)
        
        mentor_note_record = {
            "note_id": str(uuid.uuid4()),
            "user_id": uid,
            "user_email": current_user.get("email", "unknown"),
            "title": note_data.title,
            "content": note_data.content,
            "tags": note_data.tags + metadata.get("tags", []),
            "category": note_data.category or metadata.get("document_type", "general"),
            "attachment_url": note_data.attachment_url,
            "embedding": embedding,
            "ai_metadata": metadata,
            "created_timestamp": datetime.utcnow(),
            "status": "active",
            "reference_count": 0,
            "helpful_votes": 0,
            "view_count": 0
        }
        
        await db.mentor_notes.insert_one(mentor_note_record)
        
        return {
            "message": "Mentor note created successfully",
            "note_id": mentor_note_record["note_id"],
            "suggested_tags": metadata.get("tags", []),
            "category": mentor_note_record["category"]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating mentor note: {str(e)}")

@api_router.get("/knowledge/search")
async def search_knowledge_base(
    query: str,
    limit: int = 10,
    include_mentor_notes: bool = True,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Search across Community Knowledge Bank, Personal Knowledge Bank, and mentor notes"""
    try:
        uid = current_user["uid"]
        
        # Search Community Knowledge Bank (available to all users)
        community_results = await search_community_knowledge_bank(query, limit)
        
        # Search Personal Knowledge Bank (only user's own documents)
        personal_results = await search_personal_knowledge_bank(query, uid, limit)
        
        # Search mentor notes if requested
        mentor_results = []
        if include_mentor_notes:
            query_embedding = await generate_embeddings(query)
            if query_embedding:
                cursor = db.mentor_notes.find({"status": "active"})
                notes = await cursor.to_list(length=500)
                
                for note in notes:
                    if 'embedding' in note and note['embedding']:
                        note_embedding = np.array(note['embedding']).reshape(1, -1)
                        query_emb = np.array(query_embedding).reshape(1, -1)
                        similarity = cosine_similarity(query_emb, note_embedding)[0][0]
                        
                        if similarity > 0.5:  # Threshold for relevance
                            # Clean up MongoDB ObjectId for JSON serialization
                            note_clean = dict(note)
                            if '_id' in note_clean:
                                note_clean['_id'] = str(note_clean['_id'])
                            
                            mentor_results.append({
                                'type': 'mentor_note',
                                'note': note_clean,
                                'similarity_score': similarity
                            })
        
        return {
            "query": query,
            "community_results": community_results,
            "personal_results": personal_results,
            "mentor_note_results": sorted(mentor_results, key=lambda x: x['similarity_score'], reverse=True)[:5],
            "total_results": len(community_results) + len(personal_results) + len(mentor_results)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error searching knowledge base: {str(e)}")

async def search_community_knowledge_bank(query: str, limit: int = 10) -> List[Dict[str, Any]]:
    """Search Community Knowledge Bank with partner attribution"""
    try:
        query_embedding = await generate_embeddings(query)
        if not query_embedding:
            return []
            
        # Get documents from Community Knowledge Bank
        cursor = db.community_knowledge_bank.find({"status": "active"})
        documents = await cursor.to_list(length=1000)
        
        scored_docs = []
        for doc in documents:
            if 'embedding' in doc and doc['embedding']:
                doc_embedding = np.array(doc['embedding']).reshape(1, -1)
                query_emb = np.array(query_embedding).reshape(1, -1)
                similarity = cosine_similarity(query_emb, doc_embedding)[0][0]
                
                if similarity > 0.3:  # Threshold for relevance
                    # Clean up MongoDB ObjectId for JSON serialization
                    doc_clean = dict(doc)
                    if '_id' in doc_clean:
                        doc_clean['_id'] = str(doc_clean['_id'])
                    
                    # Partner content gets a boost in scoring
                    final_score = similarity * 1.2 if doc.get('partner_info') else similarity
                    
                    scored_docs.append({
                        'type': 'community_document',
                        'document': doc_clean,
                        'similarity_score': final_score,
                        'source': 'Community Knowledge Bank',
                        'company_attribution': doc.get('partner_info', {}).get('company_name', 'ONESource-ai Admin')
                    })
        
        # Sort by similarity score and return top results
        scored_docs.sort(key=lambda x: x['similarity_score'], reverse=True)
        return scored_docs[:limit]
        
    except Exception as e:
        print(f"Error in community knowledge search: {e}")
        return []

async def search_personal_knowledge_bank(query: str, user_id: str, limit: int = 10) -> List[Dict[str, Any]]:
    """Search user's Personal Knowledge Bank"""
    try:
        query_embedding = await generate_embeddings(query)
        if not query_embedding:
            return []
            
        # Get documents from user's Personal Knowledge Bank only
        cursor = db.personal_knowledge_bank.find({
            "user_id": user_id,
            "status": "active"
        })
        documents = await cursor.to_list(length=1000)
        
        scored_docs = []
        for doc in documents:
            if 'embedding' in doc and doc['embedding']:
                doc_embedding = np.array(doc['embedding']).reshape(1, -1)
                query_emb = np.array(query_embedding).reshape(1, -1)
                similarity = cosine_similarity(query_emb, doc_embedding)[0][0]
                
                if similarity > 0.3:  # Threshold for relevance
                    # Clean up MongoDB ObjectId for JSON serialization
                    doc_clean = dict(doc)
                    if '_id' in doc_clean:
                        doc_clean['_id'] = str(doc_clean['_id'])
                    
                    scored_docs.append({
                        'type': 'personal_document',
                        'document': doc_clean,
                        'similarity_score': similarity,
                        'source': 'Personal Knowledge Bank',
                        'privacy': 'Private to your account'
                    })
        
        # Sort by similarity score and return top results
        scored_docs.sort(key=lambda x: x['similarity_score'], reverse=True)
        return scored_docs[:limit]
        
    except Exception as e:
        print(f"Error in personal knowledge search: {e}")
        return []

# Enhanced Chat with Knowledge Integration
@api_router.post("/chat/ask-enhanced")
async def ask_question_enhanced(
    question_data: ChatQuestion,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Enhanced chat - UNIFIED BACKEND with knowledge bank integration"""
    try:
        # Import shared service
        from shared_chat_service import shared_chat_service
        
        uid = current_user["uid"]
        
        # Search both Community and Personal Knowledge Banks
        community_results = await search_community_knowledge_bank(question_data.question, limit=3)
        personal_results = await search_personal_knowledge_bank(question_data.question, uid, limit=2)
        
        # Build enhanced user context for shared service
        knowledge_context = []
        partner_attributions = []
        
        # Process Community Knowledge Bank results (with partner attribution)
        for result in community_results:
            if result['similarity_score'] > 0.6:  # Only high relevance
                doc = result['document']
                excerpt = doc.get('extracted_text', '')[:500]
                company_name = result.get('company_attribution', 'Community')
                
                partner_attributions.append(company_name)
                knowledge_context.append(f"From {company_name} (Community Knowledge Bank): {excerpt}")
        
        # Process Personal Knowledge Bank results (private content)
        for result in personal_results:
            if result['similarity_score'] > 0.6:  # Only high relevance
                doc = result['document']
                excerpt = doc.get('extracted_text', '')[:500]
                knowledge_context.append(f"From your personal documents: {excerpt}")
        
        # Create enhanced user context for shared service
        enhanced_context = {
            "knowledge_context": knowledge_context,
            "partner_attributions": partner_attributions,
            "community_results": len(community_results),
            "personal_results": len(personal_results)
        }
        
        # **UNIFIED BACKEND CALL - NO LEGACY LOGIC**
        response_data = shared_chat_service.get_unified_chat_response(
            question=question_data.question,
            session_id=question_data.session_id or str(uuid.uuid4()),
            user_context=enhanced_context  # Enhanced endpoint includes knowledge context
        )
        
        # Update document reference counts for Community Knowledge Bank
        for result in community_results[:3]:
            if result['similarity_score'] > 0.6:
                await db.community_knowledge_bank.update_one(
                    {"document_id": result['document']['document_id']},
                    {"$inc": {"reference_count": 1}}
                )
        
        # Update document reference counts for Personal Knowledge Bank
        for result in personal_results[:2]:
            if result['similarity_score'] > 0.6:
                await db.personal_knowledge_bank.update_one(
                    {"document_id": result['document']['document_id']},
                    {"$inc": {"reference_count": 1}}
                )
        
        # Standard conversation logging
        conversation_record = {
            "conversation_id": str(uuid.uuid4()),
            "user_id": uid,
            "session_id": response_data["session_id"],
            "question": question_data.question,
            "response": response_data["response"],  # Direct response from unified service
            "knowledge_sources_used": len(knowledge_context),
            "partner_attributions": partner_attributions,
            "community_sources": len(community_results),
            "personal_sources": len(personal_results),
            "timestamp": datetime.utcnow(),
            "tokens_used": response_data["tokens_used"]
        }
        
        await db.conversations.insert_one(conversation_record)
        
        return {
            "response": response_data["response"],  # Direct unified response (same as regular endpoint)
            "session_id": response_data["session_id"],
            "knowledge_enhanced": len(knowledge_context) > 0,
            "partner_content_used": len(partner_attributions) > 0,
            "community_sources_used": len(community_results),
            "personal_sources_used": len(personal_results),
            "tokens_used": response_data["tokens_used"],
            # Additional enhanced metadata for compatibility
            "enhanced_features": {
                "knowledge_sources": len(community_results) + len(personal_results),
                "partner_sources": partner_attributions,
                "knowledge_used": len(knowledge_context) > 0
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing enhanced question: {str(e)}")

# Feedback Routes
@api_router.post("/chat/feedback")
async def submit_feedback(
    feedback_data: ChatFeedback,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Submit feedback for a chat response"""
    try:
        uid = current_user["uid"]
        user_email = current_user.get("email", "unknown@example.com")
        
        # Prepare feedback data for storage
        feedback_record = {
            "feedback_id": str(uuid.uuid4()),
            "message_id": feedback_data.message_id,
            "user_id": uid,
            "user_email": user_email,
            "feedback_type": feedback_data.feedback_type,
            "comment": feedback_data.comment,
            "timestamp": datetime.utcnow(),
            "status": "submitted"
        }
        
        # Store feedback in MongoDB
        await db.chat_feedback.insert_one(feedback_record)
        
        # TODO: Send email notification to support team
        # This would require email service integration (SendGrid, etc.)
        print(f"📝 New feedback received:")
        print(f"   Type: {feedback_data.feedback_type}")
        print(f"   User: {user_email}")
        print(f"   Comment: {feedback_data.comment}")
        print(f"   Feedback ID: {feedback_record['feedback_id']}")
        
        return {
            "message": "Feedback submitted successfully",
            "feedback_id": feedback_record["feedback_id"],
            "note": "Your feedback has been recorded and will be reviewed by our team"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error submitting feedback: {str(e)}")

@api_router.get("/admin/feedback")
async def get_all_feedback(
    current_user: Dict[str, Any] = Depends(get_current_user),
    limit: int = 50,
    feedback_type: Optional[str] = None
):
    """Get all feedback for admin review"""
    try:
        # TODO: Add admin permission check
        # For now, allow any authenticated user to see feedback
        
        # Build query filter
        query_filter = {}
        if feedback_type and feedback_type in ["positive", "negative"]:
            query_filter["feedback_type"] = feedback_type
        
        # Get feedback from database
        cursor = db.chat_feedback.find(query_filter).sort("timestamp", -1).limit(limit)
        feedback_list = await cursor.to_list(length=limit)
        
        # Clean up MongoDB ObjectId fields
        cleaned_feedback = []
        for feedback in feedback_list:
            feedback.pop("_id", None)  # Remove MongoDB ObjectId
            cleaned_feedback.append(feedback)
        
        return {
            "feedback": cleaned_feedback,
            "total_count": len(cleaned_feedback),
            "filter_applied": feedback_type or "all"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving feedback: {str(e)}")

@api_router.post("/chat/contribution")
async def submit_contribution(
    contribution_data: KnowledgeContribution,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Submit knowledge contribution"""
    try:
        uid = current_user["uid"]
        
        # Prepare contribution data for storage
        contribution_record = {
            "contribution_id": str(uuid.uuid4()),
            "message_id": contribution_data.message_id,
            "user_id": uid,
            "user_email": current_user.get("email"),
            "user_name": current_user.get("name", current_user.get("email", "Unknown")),
            "contribution": contribution_data.contribution,
            "opt_in_credit": contribution_data.opt_in_credit,
            "timestamp": datetime.utcnow(),
            "status": "pending_review",  # pending_review, approved, rejected
            "reviewed_by": None,
            "reviewed_at": None,
            "review_notes": None
        }
        
        # Store contribution in MongoDB
        await db.knowledge_contributions.insert_one(contribution_record)
        
        return {
            "message": "Knowledge contribution submitted successfully! It will be reviewed and may be added to our knowledge base.",
            "contribution_id": contribution_record["contribution_id"]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error submitting contribution: {str(e)}")

@api_router.post("/chat/boost-response")
async def boost_response(
    boost_request: dict,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Generate a boosted response showing the next tier preview"""
    try:
        uid = current_user["uid"]
        question = boost_request.get("question")
        current_tier = boost_request.get("current_tier", "starter")
        target_tier = boost_request.get("target_tier")
        
        if not question or not target_tier:
            raise HTTPException(status_code=400, detail="Missing question or target_tier")
        
        # CRITICAL FIX: Check daily booster limit more intelligently
        today_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)
        
        booster_usage = await db.booster_usage.find_one({
            "user_id": uid,
            "date": {"$gte": today_start, "$lt": today_end}
        })
        
        # Check user's subscription tier for boost limits
        user_subscription = await firebase_service.check_user_subscription(uid)
        subscription_tier = user_subscription.get("subscription_tier", "starter")
        
        # Set daily boost limits based on subscription tier
        if subscription_tier == "starter":
            daily_limit = 1  # Free users get 1 boost per day
        elif subscription_tier in ["pro", "pro_plus"]:
            daily_limit = 10  # Pro users get more boosts per day
        else:
            daily_limit = 1  # Default to 1 for unknown tiers
        
        current_usage = booster_usage.get("usage_count", 0) if booster_usage else 0
        
        if current_usage >= daily_limit:
            limit_reset_time = (today_end).strftime("%H:%M UTC tomorrow")
            raise HTTPException(
                status_code=429, 
                detail=f"Daily booster limit reached ({current_usage}/{daily_limit}). Resets at {limit_reset_time}. Upgrade for more boosts!"
            )
        
        # Generate enhanced response based on target tier
        enhanced_system_prompt = f"""
        You are ONESource-ai, demonstrating {target_tier.upper().replace('_', '-')} tier capabilities.
        
        {target_tier.upper().replace('_', '-')} FEATURES:
        - More detailed technical analysis
        - Advanced compliance checking
        - Cross-referenced standards
        - Professional formatting with bullet points
        - Industry best practices
        - Risk assessment considerations
        {"- Specialized workflow recommendations" if target_tier == "pro_plus" else ""}
        {"- Multi-discipline coordination guidance" if target_tier == "pro_plus" else ""}
        
        Format your response with:
        - **Bold headings** for sections
        - • Bullet points for key items
        - ✅ Checkmarks for compliant items
        - ⚠️ Warnings for important considerations
        - 🏗️ Icons for construction-specific content
        
        Provide a comprehensive, professional response that clearly demonstrates the value of upgrading.
        
        Question: {question}
        """
        
        # Get AI response with enhanced prompting
        api_key = os.environ.get('OPENAI_API_KEY', '')
        if not api_key or len(api_key) < 10:
            # Enhanced mock response for booster using Enhanced Emoji Mapping
            boosted_response = f"""
# 🔧 **Technical Answer**

**Comprehensive code compliance analysis for your {question}:**

## 📊 **Code Requirements**
• Multi-standard cross-referencing (AS/NZS series)
• Primary structural requirements verified
• Fire safety protocols aligned with BCA
• Accessibility standards (DDA) compliance

## ✅ **Compliance Verification**
• Advanced risk assessment protocols completed
• Professional implementation guidelines established
• Site-specific considerations identified

# 🤓 **Mentoring Insight**

**Professional expertise considerations for this project:**

## ⚙️ **Workflow Recommendations**
• Staged construction approach recommended
• Quality control checkpoints established
• Professional certification pathways outlined

## 📄 **Documentation Needed**
• Detailed material specifications provided
• Installation methodology guidelines
• Testing and verification protocols

# 📋 **Next Steps**

## 🏛️ **Authority Requirements**
1. **Phase 1:** Initial compliance verification
2. **Phase 2:** Detailed design development  
3. **Phase 3:** Professional review and approval

## 🔄 **Alternative Solutions**
• Environmental impact assessment options
• Regulatory approval timeline variations
• Professional liability considerations

---
*This enhanced analysis demonstrates the comprehensive expertise available with {target_tier.upper().replace('_', '-')} membership.*
            """
        else:
            try:
                from openai import OpenAI
                client = OpenAI(api_key=api_key)
                
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": enhanced_system_prompt},
                        {"role": "user", "content": question}
                    ],
                    max_tokens=800,
                    temperature=0.7
                )
                
                boosted_response = response.choices[0].message.content
                
                # CRITICAL FIX: Ensure Enhanced Emoji Mapping consistency
                # Replace any incorrect emojis with the correct 🤓 emoji
                boosted_response = boosted_response.replace("🧠 **Mentoring Insight**", "🤓 **Mentoring Insight**")
                boosted_response = boosted_response.replace("💡 **Mentoring Insight**", "🤓 **Mentoring Insight**")
                boosted_response = boosted_response.replace("🧠 Mentoring Insight", "🤓 Mentoring Insight")
                boosted_response = boosted_response.replace("💡 Mentoring Insight", "🤓 Mentoring Insight")
                
            except Exception as e:
                print(f"OpenAI API error in booster: {e}")
                # Fallback to enhanced mock
                boosted_response = f"""
                **🚀 Enhanced {target_tier.upper().replace('_', '-')} Analysis**

                **Professional Assessment:**
                This comprehensive analysis demonstrates the advanced capabilities available with {target_tier.upper().replace('_', '-')} membership, including detailed compliance checking, cross-referenced standards, and professional implementation guidance.

                **Key Features Demonstrated:**
                ✅ Advanced technical analysis
                ✅ Multi-standard compliance checking  
                ✅ Professional formatting and structure
                ✅ Industry best practices integration
                ⚠️ Risk assessment and mitigation strategies

                **Value Proposition:**
                Upgrading to {target_tier.upper().replace('_', '-')} provides you with comprehensive, professional-grade responses that save time and ensure compliance across all construction disciplines.

                *Experience the full capabilities - upgrade today!*
                """
        
        # Record booster usage
        if not booster_usage:
            await db.booster_usage.insert_one({
                "user_id": uid,
                "date": today_start,
                "usage_count": 1,
                "questions_boosted": [question[:100]],
                "target_tiers": [target_tier]
            })
        else:
            await db.booster_usage.update_one(
                {"_id": booster_usage["_id"]},
                {
                    "$inc": {"usage_count": 1},
                    "$push": {
                        "questions_boosted": question[:100],
                        "target_tiers": target_tier
                    }
                }
            )
        
        return {
            "boosted_response": boosted_response,
            "target_tier": target_tier,
            "booster_used": True,
            "remaining_boosters": 0
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating boosted response: {str(e)}")

@api_router.get("/chat/history")
async def get_chat_history(
    limit: int = 50,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get user's chat history"""
    try:
        uid = current_user["uid"]
        
        # Get conversations for this user
        cursor = db.conversations.find(
            {"user_id": uid},
            sort=[("timestamp", -1)],
            limit=limit
        )
        conversations = await cursor.to_list(length=limit)
        
        # Group by session_id and create chat history entries
        session_groups = {}
        for conv in conversations:
            session_id = conv.get("session_id", "unknown")
            if session_id not in session_groups:
                session_groups[session_id] = {
                    "session_id": session_id,
                    "title": conv.get("question", "Untitled Chat")[:50] + "..." if len(conv.get("question", "")) > 50 else conv.get("question", "Untitled Chat"),
                    "timestamp": conv.get("timestamp"),
                    "messages": []
                }
            session_groups[session_id]["messages"].append({
                "question": conv.get("question"),
                "response": conv.get("formatted_response"),
                "timestamp": conv.get("timestamp")
            })
        
        # Convert to list and sort by timestamp
        chat_history = list(session_groups.values())
        chat_history.sort(key=lambda x: x.get("timestamp", datetime.min), reverse=True)
        
        return {"chat_history": chat_history}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving chat history: {str(e)}")

@api_router.get("/chat/session/{session_id}")
async def get_chat_session(
    session_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get specific chat session messages"""
    try:
        uid = current_user["uid"]
        
        # Get all messages for this session and user
        cursor = db.conversations.find(
            {"session_id": session_id, "user_id": uid},
            sort=[("timestamp", 1)]
        )
        conversations = await cursor.to_list(length=1000)
        
        messages = []
        for conv in conversations:
            # Add user message
            messages.append({
                "id": f"user_{conv['_id']}",
                "type": "user",
                "content": conv.get("question"),
                "timestamp": conv.get("timestamp").isoformat() if conv.get("timestamp") else None
            })
            
            # Add AI response
            messages.append({
                "id": f"ai_{conv['_id']}",
                "type": "ai", 
                "content": conv.get("formatted_response"),
                "timestamp": conv.get("timestamp").isoformat() if conv.get("timestamp") else None,
                "tokensUsed": conv.get("tokens_used")
            })
        
        return {"messages": messages, "session_id": session_id}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving chat session: {str(e)}")

# Developer/Admin Routes
@api_router.get("/admin/feedback")
async def get_feedback_for_review(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get feedback for developer review (basic admin function)"""
    try:
        # Basic admin check - in production, you'd want proper admin roles
        # For now, we'll just return the data for any authenticated user
        
        cursor = db.chat_feedback.find(
            {},
            sort=[("timestamp", -1)],
            limit=100
        )
        feedback = await cursor.to_list(length=100)
        
        # Clean up MongoDB ObjectId
        for item in feedback:
            item["_id"] = str(item["_id"])
        
        return {"feedback": feedback}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving feedback: {str(e)}")

@api_router.post("/admin/developer-access")
async def grant_developer_access(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Grant developer access (free consultant plan) to current user"""
    try:
        uid = current_user["uid"]
        
        # Developer access subscription data
        developer_subscription_data = {
            "subscription_tier": "consultant",
            "subscription_active": True,
            "subscription_type": "developer_access",
            "is_developer": True,
            "subscription_started_at": datetime.utcnow(),
            "payment_session_id": "developer_access_grant",
            "special_access": True
        }
        
        # Update user profile with developer access
        await firebase_service.update_subscription_status(uid, developer_subscription_data)
        
        # Log developer access grant
        access_record = {
            "access_id": str(uuid.uuid4()),
            "user_id": uid,
            "user_email": current_user.get("email"),
            "access_type": "developer",
            "granted_at": datetime.utcnow(),
            "status": "active"
        }
        
        await db.developer_access.insert_one(access_record)
        
        return {
            "message": "Developer access granted successfully! You now have unlimited consultant-level access.",
            "access_type": "developer_consultant",
            "features_unlocked": [
                "Unlimited AI queries",
                "Priority response speed",
                "Admin dashboard access",
                "Knowledge vault management",
                "Advanced features"
            ]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error granting developer access: {str(e)}")

@api_router.get("/admin/check-developer-status")
async def check_developer_status(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Check if current user has developer access"""
    try:
        uid = current_user["uid"]
        
        # Check if user has developer access in database
        developer_record = await db.developer_access.find_one(
            {"user_id": uid, "status": "active"}
        )
        
        # Also check Firebase profile
        user_profile = await firebase_service.get_user_profile(uid)
        is_developer_firebase = user_profile.get("is_developer", False) if user_profile else False
        
        return {
            "has_developer_access": developer_record is not None or is_developer_firebase,
            "access_type": "developer" if developer_record else None,
            "granted_at": developer_record.get("granted_at") if developer_record else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error checking developer status: {str(e)}")

# Voucher System Routes
@api_router.post("/admin/create-voucher")
async def create_voucher(
    voucher_data: VoucherCreate,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Create a new voucher code (admin/developer only)"""
    try:
        uid = current_user["uid"]
        
        # Check if voucher already exists
        existing_voucher = await db.vouchers.find_one({"voucher_code": voucher_data.voucher_code})
        if existing_voucher:
            raise HTTPException(status_code=400, detail="Voucher code already exists")
        
        # Create voucher record
        voucher_record = {
            "voucher_id": str(uuid.uuid4()),
            "voucher_code": voucher_data.voucher_code.upper(),
            "plan_type": voucher_data.plan_type,
            "duration_days": voucher_data.duration_days,
            "max_uses": voucher_data.max_uses,
            "current_uses": 0,
            "description": voucher_data.description,
            "created_by": uid,
            "created_at": datetime.utcnow(),
            "status": "active",
            "expires_at": None  # Vouchers don't expire, only their usage does
        }
        
        await db.vouchers.insert_one(voucher_record)
        
        return {
            "message": "Voucher created successfully",
            "voucher_code": voucher_record["voucher_code"],
            "plan_type": voucher_data.plan_type,
            "duration_days": voucher_data.duration_days,
            "max_uses": voucher_data.max_uses
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating voucher: {str(e)}")

@api_router.post("/voucher/redeem")
async def redeem_voucher(
    voucher_request: VoucherRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Redeem a voucher code"""
    try:
        uid = current_user["uid"]
        voucher_code = voucher_request.voucher_code.upper()
        
        # Find voucher
        voucher = await db.vouchers.find_one({
            "voucher_code": voucher_code,
            "status": "active"
        })
        
        if not voucher:
            raise HTTPException(status_code=404, detail="Invalid or expired voucher code")
        
        # Check if voucher has remaining uses
        if voucher["current_uses"] >= voucher["max_uses"]:
            raise HTTPException(status_code=400, detail="Voucher has no remaining uses")
        
        # Check if user already redeemed this voucher
        existing_redemption = await db.voucher_redemptions.find_one({
            "voucher_code": voucher_code,
            "user_id": uid
        })
        
        if existing_redemption:
            raise HTTPException(status_code=400, detail="You have already redeemed this voucher")
        
        # Calculate expiration date
        from datetime import timedelta
        expires_at = datetime.utcnow() + timedelta(days=voucher["duration_days"])
        
        # Create subscription data
        subscription_data = {
            "subscription_tier": voucher["plan_type"],
            "subscription_active": True,
            "subscription_type": "voucher",
            "subscription_started_at": datetime.utcnow(),
            "subscription_expires": expires_at,
            "voucher_code": voucher_code,
            "voucher_duration": voucher["duration_days"]
        }
        
        # Update user subscription
        await firebase_service.update_subscription_status(uid, subscription_data)
        
        # Record voucher redemption
        redemption_record = {
            "redemption_id": str(uuid.uuid4()),
            "voucher_code": voucher_code,
            "voucher_id": voucher["voucher_id"],
            "user_id": uid,
            "user_email": current_user.get("email"),
            "redeemed_at": datetime.utcnow(),
            "expires_at": expires_at,
            "plan_type": voucher["plan_type"],
            "duration_days": voucher["duration_days"]
        }
        
        await db.voucher_redemptions.insert_one(redemption_record)
        
        # Increment voucher usage count
        await db.vouchers.update_one(
            {"voucher_id": voucher["voucher_id"]},
            {"$inc": {"current_uses": 1}}
        )
        
        return {
            "message": f"Voucher redeemed successfully! You now have {voucher['plan_type']} access for {voucher['duration_days']} days.",
            "plan_type": voucher["plan_type"],
            "expires_at": expires_at.isoformat(),
            "duration_days": voucher["duration_days"]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error redeeming voucher: {str(e)}")

@api_router.get("/admin/vouchers")
async def list_vouchers(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """List all vouchers (admin/developer only)"""
    try:
        cursor = db.vouchers.find({}, sort=[("created_at", -1)])
        vouchers = await cursor.to_list(length=100)
        
        # Clean up MongoDB ObjectId and add redemption info
        for voucher in vouchers:
            voucher["_id"] = str(voucher["_id"])
            
            # Get redemption count
            redemption_count = await db.voucher_redemptions.count_documents({
                "voucher_code": voucher["voucher_code"]
            })
            voucher["redemption_count"] = redemption_count
        
        return {"vouchers": vouchers}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error listing vouchers: {str(e)}")

@api_router.get("/user/voucher-status")  
async def get_voucher_status(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get user's current voucher status"""
    try:
        uid = current_user["uid"]
        
        # Get active voucher redemption
        active_redemption = await db.voucher_redemptions.find_one({
            "user_id": uid,
            "expires_at": {"$gte": datetime.utcnow()}
        }, sort=[("redeemed_at", -1)])
        
        if not active_redemption:
            return {
                "has_active_voucher": False,
                "voucher_status": "none"
            }
        
        return {
            "has_active_voucher": True,
            "voucher_code": active_redemption["voucher_code"],
            "plan_type": active_redemption["plan_type"],
            "expires_at": active_redemption["expires_at"].isoformat(),
            "days_remaining": (active_redemption["expires_at"] - datetime.utcnow()).days,
            "redeemed_at": active_redemption["redeemed_at"].isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error checking voucher status: {str(e)}")

@api_router.get("/admin/contributions")
async def get_contributions_for_review(
    status: str = "pending_review",
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get knowledge contributions for developer review"""
    try:
        # Basic admin check - in production, you'd want proper admin roles
        
        query = {"status": status} if status != "all" else {}
        cursor = db.knowledge_contributions.find(
            query,
            sort=[("timestamp", -1)],
            limit=100
        )
        contributions = await cursor.to_list(length=100)
        
        # Clean up MongoDB ObjectId
        for item in contributions:
            item["_id"] = str(item["_id"])
        
        return {"contributions": contributions}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving contributions: {str(e)}")

@api_router.put("/admin/contributions/{contribution_id}")
async def review_contribution(
    contribution_id: str,
    status: str,  # approved, rejected
    review_notes: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Review and update contribution status"""
    try:
        uid = current_user["uid"]
        
        update_data = {
            "status": status,
            "reviewed_by": uid,
            "reviewed_at": datetime.utcnow(),
            "review_notes": review_notes
        }
        
        result = await db.knowledge_contributions.update_one(
            {"contribution_id": contribution_id},
            {"$set": update_data}
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Contribution not found")
        
        return {"message": f"Contribution {status} successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reviewing contribution: {str(e)}")

# Payment Routes
@api_router.post("/payment/checkout")
async def create_checkout_session(
    payment_data: PaymentRequest,
    request: Request,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user_optional)
):
    """Create Stripe checkout session"""
    try:
        user_id = current_user["uid"] if current_user else None
        
        result = await payment_service.create_checkout_session(
            payment_data.package_id,
            payment_data.origin_url,
            user_id,
            request
        )
        
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["error"])
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating checkout: {str(e)}")

@api_router.get("/payment/status/{session_id}")
async def get_checkout_status(session_id: str, request: Request):
    """Get payment status"""
    try:
        result = await payment_service.get_checkout_status(session_id, request)
        
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["error"])
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error checking payment status: {str(e)}")

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks"""
    try:
        body = await request.body()
        stripe_signature = request.headers.get("stripe-signature")
        
        if not stripe_signature:
            raise HTTPException(status_code=400, detail="Missing Stripe signature")
        
        result = await payment_service.handle_webhook(body, stripe_signature, request)
        
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result["error"])
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing webhook: {str(e)}")

@api_router.get("/pricing")
async def get_pricing():
    """Get available pricing packages"""
    from payment_service import PRICING_PACKAGES
    return {"packages": PRICING_PACKAGES}

# Weekly Business Intelligence Reporting Routes
@api_router.post("/admin/send-weekly-report")
async def send_weekly_report_endpoint(
    admin_email: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Manually trigger weekly business intelligence report"""
    try:
        # Check if user has admin access (basic implementation)
        # In production, you'd want proper admin role checking
        
        reporting_service = WeeklyReportingService()
        result = await reporting_service.send_weekly_report(admin_email)
        
        # Handle both old boolean return and new dict return
        if isinstance(result, dict):
            if result["success"]:
                return {
                    "message": "Weekly business intelligence report sent successfully",
                    "sent_to": admin_email or reporting_service.admin_email,
                    "details": result["message"]
                }
            else:
                # If SendGrid is not configured, return a different status
                if "SendGrid API key not configured" in result["message"]:
                    return {
                        "message": "Weekly report generated but email not sent",
                        "reason": "SendGrid API key not configured",
                        "suggestion": "Add valid SENDGRID_API_KEY environment variable to enable email reports",
                        "data_ready": result.get("data_collected", False),
                        "sent_to": admin_email or reporting_service.admin_email
                    }
                else:
                    return {
                        "message": "Weekly report generated but email not sent",
                        "reason": result["message"],
                        "suggestion": "Check SendGrid configuration and API key validity",
                        "data_ready": True,
                        "sent_to": admin_email or reporting_service.admin_email
                    }
        else:
            # Handle legacy boolean return
            if result:
                return {
                    "message": "Weekly business intelligence report sent successfully",
                    "sent_to": admin_email or reporting_service.admin_email
                }
            else:
                return {
                    "message": "Weekly report generated but email not sent",
                    "reason": "SendGrid not configured or invalid API key",
                    "suggestion": "Configure valid SENDGRID_API_KEY environment variable",
                    "sent_to": admin_email or reporting_service.admin_email
                }
            
    except Exception as e:
        # Check if it's a SendGrid-related error
        error_str = str(e)
        if "401" in error_str and "Unauthorized" in error_str:
            return {
                "message": "Weekly report generated but email not sent",
                "reason": "SendGrid API key invalid or unauthorized",
                "suggestion": "Check SENDGRID_API_KEY environment variable for valid API key",
                "sent_to": admin_email or "admin@onesource-ai.com",
                "error_details": "SendGrid authentication failed"
            }
        elif "SendGrid" in error_str or "email" in error_str.lower():
            return {
                "message": "Weekly report generated but email not sent",
                "reason": "Email service error",
                "suggestion": "Check SendGrid configuration",
                "sent_to": admin_email or "admin@onesource-ai.com",
                "error_details": error_str
            }
        else:
            raise HTTPException(status_code=500, detail=f"Error generating weekly report: {str(e)}")

class TestWeeklyReportRequest(BaseModel):
    admin_email: str

@api_router.post("/admin/test-weekly-report")
async def test_weekly_report_endpoint(
    request_data: TestWeeklyReportRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Test weekly business intelligence report with specific email"""
    try:
        # Check if user has admin access (basic implementation)
        
        result = await test_weekly_report(request_data.admin_email)
        
        # Handle both old boolean return and new dict return
        if isinstance(result, dict):
            if result["success"]:
                return {
                    "message": "Test weekly report sent successfully",
                    "sent_to": request_data.admin_email,
                    "details": result["message"]
                }
            else:
                # If SendGrid is not configured, return a different status
                if "SendGrid API key not configured" in result["message"]:
                    return {
                        "message": "Test report generated but email not sent",
                        "reason": "SendGrid API key not configured",
                        "suggestion": "Add valid SENDGRID_API_KEY environment variable to enable email reports",
                        "data_ready": result.get("data_collected", False),
                        "sent_to": request_data.admin_email
                    }
                else:
                    raise HTTPException(status_code=500, detail=result["message"])
        else:
            # Handle legacy boolean return
            if result:
                return {
                    "message": "Test weekly report sent successfully",
                    "sent_to": request_data.admin_email
                }
            else:
                return {
                    "message": "Test report generated but email not sent",
                    "reason": "SendGrid not configured or invalid API key",
                    "suggestion": "Configure valid SENDGRID_API_KEY environment variable",
                    "sent_to": request_data.admin_email
                }
            
    except Exception as e:
        # Check if it's a SendGrid-related error
        error_str = str(e)
        if "401" in error_str and "Unauthorized" in error_str:
            return {
                "message": "Test report generated but email not sent",
                "reason": "SendGrid API key invalid or unauthorized",
                "suggestion": "Check SENDGRID_API_KEY environment variable",
                "sent_to": request_data.admin_email,
                "error_details": "SendGrid authentication failed"
            }
        else:
            raise HTTPException(status_code=500, detail=f"Error sending test report: {str(e)}")

# User Preferences Routes
@api_router.post("/user/preferences")
async def save_user_preferences(
    preferences: UserPreferences,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Save user preferences for AI personalization"""
    try:
        uid = current_user["uid"]
        
        preference_record = {
            "user_id": uid,
            "industries": preferences.industries,
            "role": preferences.role,
            "experience_level": preferences.experience_level,
            "response_style": preferences.response_style,
            "ai_focus_areas": preferences.ai_focus_areas,
            "custom_instructions": preferences.custom_instructions,
            "updated_at": datetime.utcnow(),
            "created_at": datetime.utcnow()
        }
        
        # Upsert preferences
        await db.user_preferences.update_one(
            {"user_id": uid},
            {"$set": preference_record},
            upsert=True
        )
        
        return {"message": "Preferences saved successfully", "preferences": preference_record}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving preferences: {str(e)}")

@api_router.get("/user/preferences")
async def get_user_preferences(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get user preferences for AI personalization"""
    try:
        uid = current_user["uid"]
        
        preferences = await db.user_preferences.find_one({"user_id": uid})
        
        if not preferences:
            raise HTTPException(status_code=404, detail="No preferences found")
        
        # Clean up MongoDB ObjectId
        if "_id" in preferences:
            del preferences["_id"]
            
        return preferences
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting preferences: {str(e)}")

@api_router.get("/knowledge/personal-documents")
async def get_personal_documents(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get user's personal knowledge bank documents"""
    try:
        uid = current_user["uid"]
        
        cursor = db.personal_knowledge_bank.find(
            {"user_id": uid, "status": "active"},
            {"filename": 1, "upload_timestamp": 1, "document_id": 1, "original_size": 1, "tags": 1}
        ).sort("upload_timestamp", -1)
        
        documents = await cursor.to_list(length=100)
        
        # Clean up MongoDB ObjectIds
        for doc in documents:
            if "_id" in doc:
                del doc["_id"]
        
        return {"documents": documents, "total_count": len(documents)}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting personal documents: {str(e)}")

# Partner Management Routes
@api_router.post("/partners/register")
async def register_partner(partner_data: PartnerRegistration):
    """Register a new partner for Community Knowledge Bank"""
    try:
        if not partner_data.agreed_to_terms:
            raise HTTPException(status_code=400, detail="Must agree to Terms and Conditions")
        
        result = await partner_service.register_partner(partner_data.dict())
        
        if result["success"]:
            return {
                "message": f"Welcome {result['company_name']}! Partner registration successful.",
                "partner_id": result["partner_id"],
                "next_steps": [
                    "Check your email for welcome confirmation",
                    "Start uploading documents to Community Knowledge Bank",
                    "Your uploads will be credited to your company"
                ]
            }
        else:
            raise HTTPException(status_code=400, detail=result["message"])
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Registration failed: {str(e)}")

@api_router.get("/partners/check-status")
async def check_partner_status(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Check if current user is a registered partner"""
    try:
        uid = current_user["uid"]
        email = current_user.get("email")
        
        if not email:
            return {"is_partner": False, "message": "Email required for partner verification"}
        
        partner = await partner_service.get_partner_by_email(email)
        
        if partner:
            return {
                "is_partner": True,
                "partner_info": {
                    "company_name": partner["company_name"],
                    "partner_id": partner["partner_id"],
                    "registration_date": partner["registration_date"].isoformat(),
                    "upload_count": partner.get("upload_count", 0),
                    "status": partner.get("status", "active")
                }
            }
        else:
            return {
                "is_partner": False,
                "message": "Not registered as a partner",
                "register_url": "/partners/register"
            }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error checking partner status: {str(e)}")

@api_router.get("/admin/partners")
async def get_all_partners(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get all registered partners (admin only)"""
    try:
        # Basic admin check - in production, implement proper role checking
        partners = await partner_service.get_all_partners()
        
        return {
            "partners": partners,
            "total_count": len(partners),
            "active_count": len([p for p in partners if p.get("status") == "active"])
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting partners: {str(e)}")

@api_router.get("/admin/partners")
async def get_partner_applications(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get all partner applications for admin review"""
    try:
        # Basic admin check - in production, implement proper role checking
        partners = await partner_service.get_all_partners()
        
        return {
            "partners": partners,
            "total_count": len(partners),
            "pending_count": len([p for p in partners if p.get("status") == "pending"]),
            "approved_count": len([p for p in partners if p.get("status") == "approved"]),
            "rejected_count": len([p for p in partners if p.get("status") == "rejected"])
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting partner applications: {str(e)}")

class PartnerReviewRequest(BaseModel):
    partner_id: str
    action: str  # 'approve' or 'reject'
    admin_notes: Optional[str] = None

@api_router.post("/admin/partners/review")
async def review_partner_application(
    review_data: PartnerReviewRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Review partner application (approve/reject)"""
    try:
        # Basic admin check - in production, implement proper role checking
        admin_email = current_user.get('email', '')
        
        result = await partner_service.review_partner_application(
            partner_id=review_data.partner_id,
            action=review_data.action,
            admin_notes=review_data.admin_notes,
            reviewed_by=admin_email
        )
        
        return {
            "success": True,
            "message": f"Partner application {review_data.action}d successfully",
            "partner_id": review_data.partner_id,
            "action": review_data.action
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reviewing partner application: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
